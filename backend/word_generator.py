# backend/word_generator.py
"""
Word Document Generator voor Makelaar Contracten
Gebruikt python-docx voor simple contracts
Voor template-based: gebruik docxtpl (zie comments)
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
from typing import Dict


class ContractGenerator:
    """Genereer Word contract uit data"""
    
    def __init__(self, template_path: str = "backend/templates/template.docx"):
        self.template_path = template_path
    
    def generate_simple_contract(self, form_data: Dict, output_path: str):
        """
        Genereer een eenvoudig contract zonder template
        Perfect voor demo en testing
        """
        doc = Document()
        
        # Styling
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Titel
        title = doc.add_heading('ONDERHANDSE VERKOOPOVEREENKOMST', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Datum en plaats
        p = doc.add_paragraph()
        p.add_run(f"Opgemaakt te {form_data.get('goed_gemeente', 'België')} ")
        p.add_run(f"op {datetime.now().strftime('%d/%m/%Y')}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Lege lijn
        
        # PARTIJEN
        doc.add_heading('Tussen de partijen:', 1)
        
        # Verkoper
        doc.add_heading('VERKOPER', 2)
        p = doc.add_paragraph()
        p.add_run('Naam: ').bold = True
        p.add_run(f"{form_data.get('verkoper_voornaam', '')} {form_data.get('verkoper_naam', '')}\n")
        p.add_run('Adres: ').bold = True
        p.add_run(f"{form_data.get('verkoper_adres', '')}\n")
        if form_data.get('verkoper_email'):
            p.add_run('Email: ').bold = True
            p.add_run(f"{form_data.get('verkoper_email')}\n")
        if form_data.get('verkoper_telefoonnummer'):
            p.add_run('Telefoon: ').bold = True
            p.add_run(f"{form_data.get('verkoper_telefoonnummer')}\n")
        
        doc.add_paragraph('En', style='Intense Quote')
        
        # Koper
        doc.add_heading('KOPER', 2)
        p = doc.add_paragraph()
        p.add_run('Naam: ').bold = True
        p.add_run(f"{form_data.get('koper_voornaam', '')} {form_data.get('koper_naam', '')}\n")
        p.add_run('Adres: ').bold = True
        p.add_run(f"{form_data.get('koper_adres', '')}\n")
        if form_data.get('koper_email'):
            p.add_run('Email: ').bold = True
            p.add_run(f"{form_data.get('koper_email')}\n")
        if form_data.get('koper_telefoonnummer'):
            p.add_run('Telefoon: ').bold = True
            p.add_run(f"{form_data.get('koper_telefoonnummer')}\n")
        
        # HET GOED
        doc.add_heading('VOORWERP VAN DE OVEREENKOMST', 1)
        p = doc.add_paragraph()
        p.add_run('Adres: ').bold = True
        p.add_run(
            f"{form_data.get('goed_straat', '')} {form_data.get('goed_nummer', '')}, "
            f"{form_data.get('goed_postcode', '')} {form_data.get('goed_gemeente', '')}\n"
        )
        
        if form_data.get('goed_kadastrale_afdeling'):
            p.add_run('\nKadastrale gegevens:\n').bold = True
            p.add_run(f"• Afdeling: {form_data.get('goed_kadastrale_afdeling', '')}\n")
            p.add_run(f"• Sectie: {form_data.get('goed_kadastrale_sectie', '')}\n")
            p.add_run(f"• Nummer: {form_data.get('goed_kadastrale_nummer', '')}\n")
            if form_data.get('goed_kadastrale_oppervlakte'):
                p.add_run(f"• Oppervlakte: {form_data.get('goed_kadastrale_oppervlakte', '')}\n")
            if form_data.get('goed_kadastraal_inkomen_bedrag'):
                p.add_run(f"• Kadastraal inkomen: €{form_data.get('goed_kadastraal_inkomen_bedrag', '')}\n")
        
        # PRIJS
        doc.add_heading('PRIJS', 1)
        p = doc.add_paragraph()
        p.add_run('Koopprijs: ').bold = True
        p.add_run(f"€ {form_data.get('prijs_totaal', '0')}\n")
        p.add_run('Voorschot: ').bold = True
        p.add_run(f"€ {form_data.get('voorschot_bedrag', '0')}\n")
        
        if form_data.get('prijs_totaal') and form_data.get('voorschot_bedrag'):
            try:
                saldo = float(form_data['prijs_totaal']) - float(form_data['voorschot_bedrag'])
                p.add_run('Saldo: ').bold = True
                p.add_run(f"€ {saldo:.2f}\n")
            except:
                pass
        
        # CERTIFICATEN
        doc.add_heading('CERTIFICATEN EN ATTESTEN', 1)
        
        if form_data.get('epc_code'):
            p = doc.add_paragraph()
            p.add_run('Energieprestatiecertificaat (EPC)\n').bold = True
            p.add_run(f"• Code: {form_data.get('epc_code', '')}\n")
            p.add_run(f"• Label: {form_data.get('epc_label', '')}\n")
            p.add_run(f"• Score: {form_data.get('epc_score', '')}\n")
            p.add_run(f"• Datum: {form_data.get('epc_datum', '')}\n")
        
        if form_data.get('bodem_attest_referentie'):
            p = doc.add_paragraph()
            p.add_run('Bodemattest\n').bold = True
            p.add_run(f"• Referentie: {form_data.get('bodem_attest_referentie', '')}\n")
            p.add_run(f"• Datum: {form_data.get('bodem_attest_datum', '')}\n")
            p.add_run(f"• Inhoud: {form_data.get('bodem_attest_inhoud', '')}\n")
        
        if form_data.get('elektrische_keuring_datum'):
            p = doc.add_paragraph()
            p.add_run('Elektrische Keuring\n').bold = True
            p.add_run(f"• Datum: {form_data.get('elektrische_keuring_datum', '')}\n")
        
        if form_data.get('stedenbouw_meest_recente_bestemming'):
            p = doc.add_paragraph()
            p.add_run('Stedenbouwkundige Informatie\n').bold = True
            p.add_run(f"• Bestemming: {form_data.get('stedenbouw_meest_recente_bestemming', '')}\n")
            vergunning = 'Ja' if form_data.get('stedenbouw_vergunning_afgeleverd') else 'Nee'
            p.add_run(f"• Vergunning afgeleverd: {vergunning}\n")
        
        # ALGEMENE BEPALINGEN
        doc.add_heading('ALGEMENE BEPALINGEN', 1)
        
        p = doc.add_paragraph()
        p.add_run('Staat: ').bold = True
        p.add_run('Het goed wordt verkocht in de huidige staat, zonder waarborg voor zichtbare of verborgen gebreken.\n\n')
        
        p.add_run('Eigendomsoverdracht: ').bold = True
        p.add_run('De eigendom gaat over bij het verlijden van de authentieke akte.\n\n')
        
        p.add_run('Kosten: ').bold = True
        p.add_run('De kosten van de authentieke akte komen ten laste van de koper.\n\n')
        
        if form_data.get('notaris_verkoper') or form_data.get('notaris_koper'):
            p.add_run('Notarissen: ').bold = True
            if form_data.get('notaris_verkoper'):
                p.add_run(f"Verkoper: {form_data.get('notaris_verkoper')}. ")
            if form_data.get('notaris_koper'):
                p.add_run(f"Koper: {form_data.get('notaris_koper')}.")
            p.add_run('\n\n')
        
        # HANDTEKENINGEN
        doc.add_page_break()
        doc.add_heading('HANDTEKENINGEN', 1)
        
        p = doc.add_paragraph()
        p.add_run(f"Opgemaakt in {form_data.get('aantal_exemplaren', '3')} exemplaren te ")
        p.add_run(f"{form_data.get('goed_gemeente', 'België')} op ")
        p.add_run(f"{datetime.now().strftime('%d/%m/%Y')}.")
        
        doc.add_paragraph()
        
        # Handtekening tabel
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Headers
        table.cell(0, 0).text = 'De Verkoper'
        table.cell(0, 1).text = 'De Koper'
        
        # Ruimte voor handtekening
        table.cell(1, 0).text = '\n\n\n'
        table.cell(1, 1).text = '\n\n\n'
        
        # Namen
        table.cell(2, 0).text = f"{form_data.get('verkoper_voornaam', '')} {form_data.get('verkoper_naam', '')}"
        table.cell(2, 1).text = f"{form_data.get('koper_voornaam', '')} {form_data.get('koper_naam', '')}"
        
        # Datum
        datum = datetime.now().strftime('%d/%m/%Y')
        table.cell(3, 0).text = f"Datum: {datum}"
        table.cell(3, 1).text = f"Datum: {datum}"
        
        # Footer
        doc.add_page_break()
        p = doc.add_paragraph()
        p.add_run('Dit contract is gegenereerd door Makelaar Contract Generator.\n').italic = True
        p.add_run('Voor officieel gebruik dient dit contract door een notaris geverifieerd te worden.').italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save
        doc.save(output_path)
        return output_path
    
    def generate_from_template(self, form_data: Dict, output_path: str):
        """
        Genereer Word document uit template met docxtpl
        Voor productie met je eigen template
        
        Uncomment deze code als je docxtpl wilt gebruiken:
        
        from docxtpl import DocxTemplate
        
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Template niet gevonden: {self.template_path}")
        
        doc = DocxTemplate(self.template_path)
        
        # Prepare context met alle velden
        context = self.prepare_data(form_data)
        
        # Render template
        doc.render(context)
        
        # Save
        doc.save(output_path)
        
        return output_path
        """
        # Voor nu: gebruik simple generator
        return self.generate_simple_contract(form_data, output_path)
    
    def prepare_data(self, form_data: Dict) -> Dict:
        """
        Bereid data voor voor template
        Zet alle velden klaar en bereken afgeleide velden
        """
        context = form_data.copy()
        
        # Voeg datum toe
        context['akte_datum'] = datetime.now().strftime('%d/%m/%Y')
        context['akte_plaats'] = context.get('goed_gemeente', 'België')
        
        # Bereken saldo
        if 'prijs_totaal' in context and 'voorschot_bedrag' in context:
            try:
                prijs = float(context['prijs_totaal'])
                voorschot = float(context['voorschot_bedrag'])
                context['saldo_koopsom'] = f"{prijs - voorschot:.2f}"
            except:
                context['saldo_koopsom'] = "0"
        
        # Zorg dat alle mogelijke velden bestaan (zelfs als leeg)
        default_fields = [
            'verkoper_naam', 'verkoper_voornaam', 'verkoper_adres',
            'koper_naam', 'koper_voornaam', 'koper_adres',
            'goed_straat', 'goed_nummer', 'goed_postcode', 'goed_gemeente',
            'prijs_totaal', 'voorschot_bedrag', 'saldo_koopsom',
            'epc_code', 'epc_label', 'epc_score', 'epc_datum',
            'bodem_attest_referentie', 'bodem_attest_datum',
            'notaris_verkoper', 'notaris_koper',
            'makelaar_naam', 'aantal_exemplaren'
        ]
        
        for field in default_fields:
            if field not in context:
                context[field] = ""
        
        return context
    
    def validate_template(self, template_path: str = None) -> Dict:
        """
        Valideer of template alle vereiste velden bevat
        """
        if template_path is None:
            template_path = self.template_path
        
        if not os.path.exists(template_path):
            return {
                'valid': False,
                'error': f'Template niet gevonden: {template_path}'
            }
        
        # Voor nu: simpele check
        return {
            'valid': True,
            'message': 'Template gevonden',
            'path': template_path
        }


# Test functie
if __name__ == "__main__":
    generator = ContractGenerator()
    
    # Test data
    test_data = {
        'verkoper_naam': 'Janssens',
        'verkoper_voornaam': 'Jan',
        'verkoper_adres': 'Hoogstraat 12, 2000 Antwerpen',
        'verkoper_email': 'jan.janssens@example.com',
        
        'koper_naam': 'Peeters',
        'koper_voornaam': 'Marie',
        'koper_adres': 'Kerkstraat 45, 9000 Gent',
        'koper_email': 'marie.peeters@example.com',
        
        'goed_straat': 'Marktplein',
        'goed_nummer': '7',
        'goed_postcode': '3000',
        'goed_gemeente': 'Leuven',
        
        'goed_kadastrale_afdeling': '1',
        'goed_kadastrale_sectie': 'A',
        'goed_kadastrale_nummer': '123/02A',
        'goed_kadastrale_oppervlakte': '450 m²',
        
        'prijs_totaal': '350000',
        'voorschot_bedrag': '35000',
        
        'epc_code': 'EPC-2024-1234-5678',
        'epc_label': 'C',
        'epc_score': '250 kWh/m²',
        'epc_datum': '2024-03-15',
        
        'bodem_attest_referentie': 'OVAM-2024-001234',
        'bodem_attest_datum': '2024-02-10',
        
        'stedenbouw_meest_recente_bestemming': 'Woongebied',
        'stedenbouw_vergunning_afgeleverd': True
    }
    
    # Genereer test contract
    output = "test_contract.docx"
    generator.generate_simple_contract(test_data, output)
    print(f"✅ Test contract gegenereerd: {output}")
